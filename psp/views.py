from decimal import Decimal, InvalidOperation
from io import BytesIO
import os

from django.conf import settings
from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.decorators import login_required
from django.core.exceptions import PermissionDenied
from django.db import transaction
from django.db.models import Q
from django.db.models.deletion import ProtectedError
from django.http import FileResponse, HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse_lazy
from django.utils import timezone
from django.views import View
from django.views.generic import CreateView, DeleteView, DetailView, UpdateView
from openpyxl import Workbook, load_workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT, TA_JUSTIFY
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image

from core.access import get_user_unit_kerja, is_biro_umum_user, is_global_bmn_scope_user, require_user_unit_or_all, scope_queryset_by_user
from core.roles import is_sekretaris_jenderal
from core.listing import SearchListMixin
from core.export_utils import apply_search_filter, export_queryset
from .forms import ImportBarangPSPForm, PermohonanPSPBMNForm
from .models import BarangPSP, FotoBarangPSP, PermohonanPSPBMN


class PermohonanPSPAccessMixin(LoginRequiredMixin):
    def get_scoped_queryset(self):
        qs = PermohonanPSPBMN.objects.select_related(
            'unit_kerja', 'pemohon', 'dibuat_oleh', 'diverifikasi_oleh', 'disetujui_sekjen_oleh'
        ).prefetch_related('detail_barang')
        # Scope PSP:
        # - Sekretaris Jenderal/Admin System melihat semua PSP seluruh unit kerja,
        #   karena PSP ditetapkan/ditindaklanjuti pada level Sekjen.
        # - BMN Sekretariat kantor pusat Eselon I: Sekretariat + Direktorat di bawah Eselon I-nya.
        # - BMN Sentra/Balai: hanya Sentra/Balai miliknya sendiri.
        if is_sekretaris_jenderal(self.request.user) or self.request.user.is_superuser:
            return qs
        return scope_queryset_by_user(qs, self.request.user, 'psp')

    def get_queryset(self):
        return self.get_scoped_queryset()

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['user'] = self.request.user
        return kwargs


class SafeDeleteMixin:
    template_name = 'includes/confirm_delete.html'
    success_message = 'Data berhasil dihapus.'
    protected_message = 'Data tidak dapat dihapus karena masih digunakan oleh data lain.'

    def form_valid(self, form):
        try:
            messages.success(self.request, self.success_message)
            return super().form_valid(form)
        except ProtectedError:
            messages.error(self.request, self.protected_message)
            return redirect(self.get_success_url())



class PermohonanPSPListView(PermohonanPSPAccessMixin, SearchListMixin):
    """Daftar permohonan PSP untuk pemohon/unit kerja.

    Fitur verifikasi Biro Umum dan persetujuan Sekjen sengaja dipisahkan
    agar alur bisnis menjadi jelas:
    Unit Kerja -> Verifikasi Biro Umum -> Penetapan Sekjen.
    """
    model = PermohonanPSPBMN
    template_name = 'psp/list.html'
    select_related = ['unit_kerja', 'pemohon']
    mode = 'permohonan'
    page_title = 'Permohonan PSP BMN'
    empty_message = 'Belum ada permohonan PSP BMN.'
    search_fields = [
        ('nomor_permohonan', 'Nomor Permohonan'),
        ('nomor_nota_permohonan_psp', 'Nomor Nota ke Sekjen'),
        ('nomor_tiket_siman', 'Tiket SIMAN'),
        ('unit_kerja__nama_unit', 'Unit Kerja'),
        ('pemohon__nama', 'Nama Pemohon'),
        ('pemohon__nip', 'NIP Pemohon'),
        ('judul_paket', 'Judul Paket'),
        ('jenis_barang', 'Jenis Barang'),
        ('status', 'Status'),
        ('nomor_sk_psp', 'Nomor SK PSP'),
    ]

    def get_base_queryset_for_mode(self):
        if self.mode == 'verifikasi':
            if not (is_biro_umum_user(self.request.user) or self.request.user.is_superuser):
                raise PermissionDenied('Hanya role Biro Umum yang dapat mengakses Verifikasi Usulan PSP.')
            # Biro Umum memverifikasi semua usulan dari unit kerja, pusat maupun sentra/balai.
            return PermohonanPSPBMN.objects.select_related('unit_kerja', 'pemohon').prefetch_related('detail_barang').exclude(status='DRAFT')
        if self.mode == 'persetujuan_sekjen':
            if not (is_sekretaris_jenderal(self.request.user) or self.request.user.is_superuser):
                raise PermissionDenied('Hanya role Sekjen yang dapat mengakses Persetujuan PSP.')
            # Sekjen melihat semua PSP lintas unit kerja untuk penetapan.
            return PermohonanPSPBMN.objects.select_related('unit_kerja', 'pemohon').prefetch_related('detail_barang').exclude(status='DRAFT')
        return self.get_scoped_queryset()

    def get_queryset(self):
        qs = self.get_base_queryset_for_mode()
        q = (self.request.GET.get('q') or '').strip()
        selected_field = (self.request.GET.get('search_field') or 'ALL').strip()
        if q and self.search_fields:
            available_fields = [field for field, _label in self.search_fields]
            fields_to_search = available_fields if selected_field == 'ALL' or selected_field not in available_fields else [selected_field]
            query = Q()
            for field in fields_to_search:
                query |= Q(**{f'{field}__icontains': q})
            qs = qs.filter(query)
        return qs

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx.update({
            'mode': self.mode,
            'page_title': self.page_title,
            'empty_message': self.empty_message,
            'show_create_buttons': self.mode == 'permohonan' and not is_sekretaris_jenderal(self.request.user),
            'show_unit_edit_actions': self.mode == 'permohonan' and not is_sekretaris_jenderal(self.request.user),
            'show_verifikasi_actions': self.mode == 'verifikasi',
            'show_persetujuan_actions': self.mode == 'persetujuan_sekjen',
        })
        return ctx


class VerifikasiPSPListView(PermohonanPSPListView):
    mode = 'verifikasi'
    page_title = 'Verifikasi Usulan PSP BMN - Biro Umum'
    empty_message = 'Belum ada usulan PSP yang perlu diverifikasi Biro Umum.'


class PersetujuanSekjenPSPListView(PermohonanPSPListView):
    mode = 'persetujuan_sekjen'
    page_title = 'Persetujuan/Penetapan PSP BMN - Sekjen'
    empty_message = 'Belum ada usulan PSP yang diajukan ke Sekjen.'

def _save_foto_barang_files(request, permohonan):
    for foto in request.FILES.getlist('foto_barang_files'):
        FotoBarangPSP.objects.create(
            permohonan=permohonan,
            foto=foto,
            diupload_oleh=request.user if request.user.is_authenticated else None,
        )


class PermohonanPSPCreateView(PermohonanPSPAccessMixin, CreateView):
    model = PermohonanPSPBMN
    form_class = PermohonanPSPBMNForm
    template_name = 'psp/form.html'
    success_url = reverse_lazy('psp:list')

    def form_valid(self, form):
        user = self.request.user
        if not (is_global_bmn_scope_user(user) or is_sekretaris_jenderal(user)):
            unit = get_user_unit_kerja(user)
            if not unit:
                raise PermissionDenied('User belum memiliki Unit Kerja/Satker.')
            form.instance.unit_kerja = unit
            form.instance.status = 'DIAJUKAN'
        elif not form.instance.status:
            form.instance.status = 'DIAJUKAN'
        form.instance.dibuat_oleh = user
        form.instance.diperbarui_oleh = user
        form.instance.nama_barang = form.instance.nama_barang or form.instance.judul_paket or form.instance.get_jenis_barang_display()
        response = super().form_valid(form)
        _save_foto_barang_files(self.request, self.object)
        self.object.refresh_rekap_barang(commit=True)
        messages.success(self.request, 'Permohonan PSP BMN berhasil disimpan. Nomor surat otomatis sudah dibuat bila kolom nomor dikosongkan.')
        return response


class PermohonanPSPUpdateView(PermohonanPSPAccessMixin, UpdateView):
    model = PermohonanPSPBMN
    form_class = PermohonanPSPBMNForm
    template_name = 'psp/form.html'
    success_url = reverse_lazy('psp:list')

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        if is_sekretaris_jenderal(request.user) and not request.user.is_superuser:
            raise PermissionDenied('Role Sekjen hanya dapat melihat/detail Permohonan PSP. Edit umum tidak diperbolehkan.')
        # Role BMN tetap diizinkan edit usulan PSP melalui menu Permohonan PSP BMN
        # sepanjang objek berada dalam scope kewenangannya. Role Sekjen tidak boleh
        # memakai form edit umum.
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        user = self.request.user
        form.instance.diperbarui_oleh = user
        if is_global_bmn_scope_user(user) or is_sekretaris_jenderal(user):
            if form.instance.status in ['DIVERIFIKASI_BIRO', 'SIAP_DIAJUKAN_SEKJEN', 'DIAJUKAN_SEKJEN', 'DISETUJUI_SEKJEN', 'DIAJUKAN_BIRO_HUKUM', 'SK_TERBIT', 'SELESAI'] and not form.instance.tanggal_verifikasi:
                form.instance.tanggal_verifikasi = timezone.now().date()
                form.instance.diverifikasi_oleh = user
            if form.instance.status in ['DISETUJUI_SEKJEN', 'DIAJUKAN_BIRO_HUKUM', 'SK_TERBIT', 'SELESAI'] and not form.instance.tanggal_persetujuan_sekjen:
                form.instance.tanggal_persetujuan_sekjen = timezone.now().date()
                form.instance.disetujui_sekjen_oleh = user
            if form.instance.status_tte == 'SUDAH_TTE' and not form.instance.tanggal_tte:
                form.instance.tanggal_tte = timezone.now()
        else:
            form.instance.status = 'DIAJUKAN'
        response = super().form_valid(form)
        _save_foto_barang_files(self.request, self.object)
        self.object.refresh_rekap_barang(commit=True)
        messages.success(self.request, 'Permohonan PSP BMN berhasil diperbarui.')
        return response


class PermohonanPSPDetailView(PermohonanPSPAccessMixin, DetailView):
    model = PermohonanPSPBMN
    template_name = 'psp/detail.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['foto_barang_list'] = self.object.foto_barang_list.all()
        ctx['detail_barang'] = self.object.detail_barang.all()[:100]
        ctx['detail_barang_count'] = self.object.detail_barang.count()
        ctx['dokumen_checklist'] = self.object.kelengkapan_dokumen
        ctx['can_verifikasi_psp'] = (is_biro_umum_user(self.request.user) or self.request.user.is_superuser) and self.object.status in ['DIAJUKAN', 'DIVERIFIKASI_BIRO', 'PERLU_PERBAIKAN']
        ctx['can_penetapan_psp'] = (is_sekretaris_jenderal(self.request.user) or self.request.user.is_superuser) and self.object.status in ['DIAJUKAN_SEKJEN', 'SIAP_DIAJUKAN_SEKJEN', 'DISETUJUI_SEKJEN', 'SK_TERBIT']
        ctx['back_url_name'] = self.request.GET.get('from') or ('psp:persetujuan_sekjen' if is_sekretaris_jenderal(self.request.user) else ('psp:verifikasi' if is_biro_umum_user(self.request.user) else 'psp:list'))
        return ctx


class PermohonanPSPDeleteView(PermohonanPSPAccessMixin, SafeDeleteMixin, DeleteView):
    model = PermohonanPSPBMN
    success_url = reverse_lazy('psp:list')
    success_message = 'Permohonan PSP BMN berhasil dihapus.'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        if is_sekretaris_jenderal(request.user) and not request.user.is_superuser:
            raise PermissionDenied('Role Sekjen hanya dapat melihat/detail Permohonan PSP. Hapus tidak diperbolehkan.')
        if not is_global_bmn_scope_user(request.user) and self.object.status not in ['DRAFT', 'DIAJUKAN', 'PERLU_PERBAIKAN']:
            raise PermissionDenied('Usulan PSP yang sudah diproses Biro Umum tidak dapat dihapus oleh unit kerja.')
        return super().dispatch(request, *args, **kwargs)


def _normalize_header(value):
    if value is None:
        return ''
    return str(value).strip().lower().replace(' ', '_').replace('-', '_')


def _decimal(value):
    if value in [None, '']:
        return Decimal('0')
    if isinstance(value, (int, float, Decimal)):
        return Decimal(str(value))
    cleaned = str(value).strip().replace('Rp', '').replace('rp', '').replace(' ', '')
    cleaned = cleaned.replace('.', '').replace(',', '.') if ',' in cleaned else cleaned.replace(',', '')
    try:
        return Decimal(cleaned)
    except InvalidOperation:
        return Decimal('0')


def _condition_code(value):
    text = (str(value or 'Baik').strip().upper().replace(' ', '_'))
    mapping = {
        'BAIK': 'BAIK',
        'B': 'BAIK',
        'RUSAK_RINGAN': 'RUSAK_RINGAN',
        'RR': 'RUSAK_RINGAN',
        'RUSAK_BERAT': 'RUSAK_BERAT',
        'RB': 'RUSAK_BERAT',
    }
    return mapping.get(text, 'LAINNYA')



class ProsesPSPView(PermohonanPSPAccessMixin, View):
    """Aksi workflow PSP tanpa memakai form edit umum.

    Alur: Unit Kerja mengajukan -> Biro Umum verifikasi/teruskan -> Sekjen menetapkan SK.
    """
    def post(self, request, *args, **kwargs):
        obj = get_object_or_404(PermohonanPSPBMN.objects.select_related('unit_kerja', 'pemohon'), pk=kwargs['pk'])
        action = request.POST.get('action')
        catatan = (request.POST.get('catatan') or '').strip()
        today = timezone.now().date()

        if action in ['verifikasi_biro', 'kembalikan_biro', 'teruskan_sekjen']:
            if not (is_biro_umum_user(request.user) or request.user.is_superuser):
                raise PermissionDenied('Hanya role Biro Umum yang dapat memverifikasi usulan PSP.')

            uploaded_dokumen_psp = request.FILES.get('dokumen_permohonan_psp')
            if uploaded_dokumen_psp:
                if not uploaded_dokumen_psp.name.lower().endswith('.pdf'):
                    messages.error(request, 'Dokumen PSP SIKARIS final/gabungan wajib berformat PDF.')
                    return redirect('psp:detail', pk=obj.pk)
                obj.dokumen_permohonan_psp = uploaded_dokumen_psp

            obj.diverifikasi_oleh = request.user
            obj.tanggal_verifikasi = obj.tanggal_verifikasi or today
            if catatan:
                obj.catatan_biro_umum = catatan
            if action == 'verifikasi_biro':
                obj.status = 'DIVERIFIKASI_BIRO'
                messages.success(request, 'Usulan PSP berhasil diverifikasi Biro Umum.')
            elif action == 'kembalikan_biro':
                obj.status = 'PERLU_PERBAIKAN'
                messages.success(request, 'Usulan PSP dikembalikan ke unit kerja untuk perbaikan.')
            else:
                if not obj.dokumen_permohonan_psp:
                    messages.error(request, 'Sebelum diteruskan ke Sekjen, upload Dokumen PSP SIKARIS final/gabungan dalam format PDF. PDF dapat bertanda tangan elektronik BSrE atau tanda tangan manual, serta dapat memakai meterai biasa atau e-Meterai sesuai kebutuhan.')
                    return redirect('psp:detail', pk=obj.pk)
                obj.status = 'DIAJUKAN_SEKJEN'
                obj.tanggal_nota_permohonan_psp = obj.tanggal_nota_permohonan_psp or today
                messages.success(request, 'Usulan PSP berhasil diteruskan ke Sekjen untuk penetapan.')
            obj.diperbarui_oleh = request.user
            obj.save()
            return redirect('psp:detail', pk=obj.pk)

        if action in ['tetapkan_sekjen', 'tolak_sekjen']:
            if not (is_sekretaris_jenderal(request.user) or request.user.is_superuser):
                raise PermissionDenied('Hanya role Sekjen yang dapat menetapkan PSP.')
            if action == 'tetapkan_sekjen':
                obj.status = 'SK_TERBIT'
                obj.disetujui_sekjen_oleh = request.user
                obj.tanggal_persetujuan_sekjen = obj.tanggal_persetujuan_sekjen or today
                obj.tanggal_sk_psp = obj.tanggal_sk_psp or today
                if request.FILES.get('sk_penetapan_psp'):
                    obj.sk_penetapan_psp = request.FILES['sk_penetapan_psp']
                if request.FILES.get('file_setelah_tte'):
                    obj.file_setelah_tte = request.FILES['file_setelah_tte']
                    obj.status_tte = 'SUDAH_TTE'
                    obj.tanggal_tte = timezone.now()
                messages.success(request, 'PSP BMN berhasil ditetapkan oleh Sekjen.')
            else:
                obj.status = 'DITOLAK'
                if catatan:
                    obj.catatan_unit = catatan
                messages.success(request, 'Usulan PSP ditolak oleh Sekjen.')
            obj.diperbarui_oleh = request.user
            obj.save()
            return redirect('psp:detail', pk=obj.pk)

        raise PermissionDenied('Aksi proses PSP tidak dikenal.')

class ImportBarangPSPView(PermohonanPSPAccessMixin, View):
    template_name = 'psp/import_barang.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(self.get_scoped_queryset(), pk=kwargs['pk'])
        if is_sekretaris_jenderal(request.user) and not request.user.is_superuser:
            raise PermissionDenied('Role Sekjen hanya dapat melihat/detail Permohonan PSP. Import barang tidak diperbolehkan.')
        # Role BMN tetap diizinkan import barang PSP melalui menu Permohonan PSP BMN
        # sepanjang objek berada dalam scope kewenangannya. Role Sekjen tidak boleh import.
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        return render(request, self.template_name, {'object': self.object, 'form': ImportBarangPSPForm()})

    @transaction.atomic
    def post(self, request, *args, **kwargs):
        form = ImportBarangPSPForm(request.POST, request.FILES)
        if not form.is_valid():
            return render(request, self.template_name, {'object': self.object, 'form': form})

        if form.cleaned_data.get('replace_existing'):
            self.object.detail_barang.all().delete()

        wb = load_workbook(form.cleaned_data['file_excel'], data_only=True)
        ws = wb.active
        header = [_normalize_header(c.value) for c in ws[1]]
        required = ['kode_barang', 'nup', 'nama_barang', 'nilai_perolehan']
        missing = [h for h in required if h not in header]
        if missing:
            messages.error(request, 'Kolom wajib tidak ditemukan: ' + ', '.join(missing))
            return render(request, self.template_name, {'object': self.object, 'form': form})

        idx = {name: header.index(name) for name in header if name}
        created = 0
        errors = []
        for row_no, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(row):
                continue
            def pick(*names):
                for name in names:
                    if name in idx and idx[name] < len(row):
                        return row[idx[name]]
                return None
            try:
                kode_barang = str(pick('kode_barang') or '').strip()
                nup = str(pick('nup') or '').strip()
                nama_barang = str(pick('nama_barang') or '').strip()
                if not kode_barang or not nup or not nama_barang:
                    raise ValueError('kode_barang, nup, dan nama_barang wajib diisi')
                kuantitas = int(pick('kuantitas', 'qty', 'jumlah') or 1)
                nilai = _decimal(pick('nilai_perolehan', 'nilai'))
                defaults = {
                    'nomor_urut': int(pick('no', 'nomor_urut') or row_no - 1),
                    'kode_satuan_kerja': pick('kode_satuan_kerja', 'kode_satker') or self.object.kode_satuan_kerja,
                    'nama_satuan_kerja': pick('nama_satuan_kerja', 'nama_satker') or self.object.nama_satuan_kerja,
                    'nama_barang': nama_barang,
                    'tipe_barang': pick('tipe_barang', 'tipe'),
                    'tahun_perolehan': str(pick('tahun_perolehan') or '').strip(),
                    'kuantitas': kuantitas,
                    'nilai_perolehan': nilai,
                    'kondisi_barang': _condition_code(pick('kondisi_barang', 'kondisi')),
                    'keterangan': pick('keterangan'),
                }
                BarangPSP.objects.update_or_create(
                    permohonan=self.object,
                    kode_barang=kode_barang,
                    nup=nup,
                    defaults=defaults,
                )
                created += 1
            except Exception as exc:
                errors.append(f'Baris {row_no}: {exc}')
                if len(errors) >= 10:
                    break

        self.object.refresh_rekap_barang(commit=True)
        if errors:
            messages.warning(request, f'Import selesai sebagian. Berhasil {created} baris. Error: ' + '; '.join(errors))
        else:
            messages.success(request, f'Import detail barang PSP berhasil: {created} baris. Rekap total sudah diperbarui otomatis.')
        return redirect('psp:detail', pk=self.object.pk)



# =====================================================================
# Dokumen konsep PSP hasil generate aplikasi SIKARIS
# =====================================================================
def _rupiah_plain(value):
    try:
        return 'Rp ' + f'{float(value or 0):,.0f}'.replace(',', '.')
    except Exception:
        return 'Rp 0'


def _pdf_response(story, filename, pagesize=A4):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=pagesize, leftMargin=1.8*cm, rightMargin=1.8*cm, topMargin=1.5*cm, bottomMargin=1.5*cm)
    doc.build(story)
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=False, filename=filename)



def _psp_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name='KopTitle', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica-Bold', fontSize=14, leading=16, spaceAfter=0
    ))
    styles.add(ParagraphStyle(
        name='KopUnit', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica-Bold', fontSize=10, leading=12, spaceAfter=0
    ))
    styles.add(ParagraphStyle(
        name='KopAddress', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica', fontSize=7, leading=9, spaceAfter=2
    ))
    styles.add(ParagraphStyle(
        name='DocTitle', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica-Bold', fontSize=13, leading=15, spaceBefore=8, spaceAfter=1
    ))
    styles.add(ParagraphStyle(
        name='DocNumber', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica-Bold', fontSize=9, leading=11, spaceAfter=8
    ))
    styles.add(ParagraphStyle(
        name='BodyJustify', parent=styles['Normal'], alignment=TA_JUSTIFY,
        fontName='Helvetica', fontSize=8.8, leading=11
    ))
    styles.add(ParagraphStyle(
        name='BodySmall', parent=styles['Normal'], alignment=TA_LEFT,
        fontName='Helvetica', fontSize=8.5, leading=10.5
    ))
    styles.add(ParagraphStyle(
        name='Sign', parent=styles['Normal'], alignment=TA_CENTER,
        fontName='Helvetica', fontSize=8.5, leading=10.5
    ))
    styles.add(ParagraphStyle(
        name='ItalicSmall', parent=styles['Normal'], alignment=TA_LEFT,
        fontName='Helvetica-Oblique', fontSize=8.2, leading=10
    ))
    return styles


def _kop_psp(story, styles):
    """Kop dokumen PSP mengikuti contoh Biro Umum dengan logo Kemensos di kiri atas."""
    logo_candidates = [
        os.path.join(settings.BASE_DIR, 'static', 'img', 'logo-kemensos.png'),
        os.path.join(settings.BASE_DIR, 'static', 'img', 'logo-kemensos.svg'),
        os.path.join(settings.BASE_DIR, 'logo-kemensos.png'),
        os.path.join(settings.BASE_DIR, 'logo-kemensos.svg'),
    ]
    logo_path = next((path for path in logo_candidates if os.path.exists(path)), None)

    header_text = [
        Paragraph('KEMENTERIAN SOSIAL REPUBLIK INDONESIA', styles['KopTitle']),
        Paragraph('SEKRETARIAT JENDERAL', styles['KopUnit']),
        Paragraph('BIRO UMUM', styles['KopUnit']),
        Paragraph('JALAN SALEMBA RAYA NOMOR 28 JAKARTA PUSAT TELEPON : 021 - 3103591 EXT. : 2109/3103743 LAMAN: http://www.kemensos.go.id', styles['KopAddress']),
    ]

    if logo_path and logo_path.lower().endswith(('.png', '.jpg', '.jpeg')):
        logo = Image(logo_path, width=2.0*cm, height=2.0*cm)
        kop = Table([[logo, header_text]], colWidths=[2.3*cm, 15.1*cm])
    else:
        # Fallback aman jika hanya tersedia SVG atau logo belum tersedia.
        kop = Table([['', header_text]], colWidths=[2.3*cm, 15.1*cm])

    kop.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))
    story.append(kop)
    line = Table([['']], colWidths=[17.4*cm], rowHeights=[0.08*cm])
    line.setStyle(TableStyle([('LINEBELOW', (0,0), (-1,-1), 1.1, colors.black)]))
    story.append(line)
    story.append(Spacer(1, 0.35*cm))

def _psp_doc_date(value):
    value = value or timezone.now().date()
    bulan = ['Januari','Februari','Maret','April','Mei','Juni','Juli','Agustus','September','Oktober','November','Desember']
    try:
        return f'{value.day} {bulan[value.month-1]} {value.year}'
    except Exception:
        return str(value)


def _pejabat_psp(obj):
    """Ambil jabatan penandatangan PSP untuk dokumen Word.

    Nama dan NIP sengaja dikosongkan apabila belum diisi eksplisit pada data PSP,
    agar dokumen Word tidak keliru mengambil pegawai dummy sebagai Kepala Biro Umum.
    Pengguna dapat mengisi manual di file Word sebelum proses TTE BSrE.
    """
    nama = (getattr(obj, 'pejabat_tte', None) or '').strip()
    nip = (getattr(obj, 'nip_pejabat_tte', None) or '').strip()
    jabatan = 'Kepala Biro Umum'

    # Jika nama/NIP terlihat berasal dari data dummy, kosongkan agar tidak salah pejabat.
    dummy_tokens = ['BIRO UMUM 0', 'DUMMY', 'PEGAWAI']
    if any(token in nama.upper() for token in dummy_tokens):
        nama = ''
        nip = ''

    return nama, nip, jabatan

def _psp_summary(obj):
    satuan = obj.nama_satuan_kerja or (getattr(obj.unit_kerja, 'nama_unit', None) if obj.unit_kerja else '-') or '-'
    jumlah = obj.jumlah_barang or 0
    nilai = obj.total_nilai_barang or obj.nilai_psp or 0
    judul = obj.judul_paket or obj.nama_barang or obj.get_jenis_barang_display()
    tiket = obj.nomor_tiket_siman or '-'
    return satuan, jumlah, nilai, judul, tiket


def _signature_block(story, styles, jabatan='Kepala Biro Umum', nama='-', nip='-', tanggal=None, meterai=False):
    story.append(Spacer(1, 0.35*cm))
    story.append(Table([
        ['', Paragraph(f'Jakarta, {_psp_doc_date(tanggal)}<br/>{jabatan}', styles['Sign'])],
        ['', Paragraph('<br/><br/><br/>' + ('[Area e-Meterai elektronik]  ' if meterai else '') + '<br/>', styles['Sign'])],
        ['', Paragraph(f'<b>{nama}</b><br/>NIP. {nip}', styles['Sign'])],
    ], colWidths=[9.0*cm, 7.8*cm]))


def _docx_set_table_borders_none(table):
    """Hilangkan border tabel agar tabel hanya dipakai untuk tata letak Word."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = 'w:' + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'nil')


def _docx_set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def _docx_set_font(run, size=10, bold=False):
    from docx.shared import Pt
    from docx.oxml.ns import qn
    run.bold = bold
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    try:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    except Exception:
        pass


def _docx_add_kop_psp(doc):
    """Kop Word PSP dibuat ringkas mengikuti contoh PDF: logo kiri, teks kop di tengah, garis bawah."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    section = doc.sections[0]
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.65)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(2.35)
    table.columns[1].width = Cm(14.1)
    _docx_set_table_borders_none(table)
    cell_logo = table.cell(0, 0)
    cell_text = table.cell(0, 1)
    _docx_set_cell_margins(cell_logo, 0, 0, 0, 0)
    _docx_set_cell_margins(cell_text, 0, 0, 0, 0)

    # bersihkan paragraf awal
    for cell in (cell_logo, cell_text):
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    logo_candidates = [
        os.path.join(settings.BASE_DIR, 'static', 'img', 'logo-kemensos.png'),
        os.path.join(settings.BASE_DIR, 'static', 'img', 'logo_kemensos.png'),
        os.path.join(settings.BASE_DIR, 'static', 'img', 'logo.png'),
        os.path.join(settings.BASE_DIR, 'logo-kemensos.png'),
    ]
    logo_path = next((path for path in logo_candidates if os.path.exists(path)), None)
    p_logo = cell_logo.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if logo_path:
        try:
            p_logo.add_run().add_picture(logo_path, width=Cm(1.85))
        except Exception:
            pass

    # gunakan paragraf pertama cell_text agar tidak ada baris kosong di atas kop
    texts = [
        ('KEMENTERIAN SOSIAL REPUBLIK INDONESIA', 13.5, True),
        ('SEKRETARIAT JENDERAL', 11, True),
        ('BIRO UMUM', 11, True),
        ('JALAN SALEMBA RAYA NOMOR 28 JAKARTA PUSAT TELEPON : 021 - 3103591 EXT. : 2109/3103743 LAMAN: http://www.kemensos.go.id', 7.5, False),
    ]
    for idx, (text, size, bold) in enumerate(texts):
        p = cell_text.paragraphs[0] if idx == 0 else cell_text.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0 if idx < 3 else 2)
        r = p.add_run(text)
        _docx_set_font(r, size=size, bold=bold)

    # garis bawah kop
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_p(doc, text='', bold=False, align=None, size=10, space_after=3, style=None, left_indent=None, first_line_indent=None):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, Cm
    p = doc.add_paragraph(style=style)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    r = p.add_run(text)
    _docx_set_font(r, size=size, bold=bold)
    return p


def _docx_meta_table(doc, rows):
    from docx.shared import Cm, Pt
    table = doc.add_table(rows=0, cols=3)
    table.autofit = False
    table.columns[0].width = Cm(2.35)
    table.columns[1].width = Cm(0.28)
    table.columns[2].width = Cm(13.6)
    _docx_set_table_borders_none(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = ':'
        cells[2].text = str(value)
        for c in cells:
            _docx_set_cell_margins(c, 0, 0, 0, 0)
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    _docx_set_font(r, size=10)
    return table


def _docx_identitas_table(doc, nama, nip, jabatan):
    _docx_meta_table(doc, [('Nama', nama or ''), ('NIP', nip or ''), ('Jabatan', jabatan or '')])


def _docx_signature(doc, jabatan='Kepala Biro Umum', nama='', nip='', tanggal=None, meterai=False):
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(9.2)
    table.columns[1].width = Cm(7.0)
    _docx_set_table_borders_none(table)
    cell = table.cell(0, 1)
    _docx_set_cell_margins(cell, 0, 0, 0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    for text in [f'Jakarta, {_psp_doc_date(tanggal)}\n', f'{jabatan}\n']:
        r = p.add_run(text)
        _docx_set_font(r, size=10)
    if meterai:
        r = p.add_run('\n[Area e-Meterai elektronik]\n')
        _docx_set_font(r, size=10, bold=True)
    r = p.add_run('\n\n')
    _docx_set_font(r, size=10)
    r = p.add_run((nama or '____________________________') + '\n')
    _docx_set_font(r, size=10, bold=bool(nama))
    r = p.add_run('NIP. ' + (nip or '________________'))
    _docx_set_font(r, size=10)


def _docx_response(doc, filename):
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    response = FileResponse(
        buffer,
        as_attachment=True,
        filename=filename,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    return response


class GenerateDokumenPSPView(PermohonanPSPAccessMixin, View):
    """Generate konsep dokumen PSP dalam format Word (.docx) agar dapat diedit.

    jenis:
    - nota: Nota Dinas permohonan PSP dari Biro Umum ke Sekjen.
    - keterangan: Surat Keterangan kebenaran dokumen digital.
    - pernyataan: Surat Pernyataan Formil dan Materiil dengan area e-Meterai.
    - sk: Konsep SK Penetapan PSP untuk Sekjen tanpa lampiran daftar barang.
    """
    def get(self, request, *args, **kwargs):
        obj = get_object_or_404(self.get_scoped_queryset(), pk=kwargs['pk'])
        jenis = kwargs.get('jenis')
        try:
            from docx import Document
            from docx.shared import Pt
        except Exception:
            messages.error(request, 'Library python-docx belum terpasang. Jalankan: pip install python-docx')
            return redirect('psp:detail', pk=obj.pk)

        doc = Document()
        normal = doc.styles['Normal']
        normal.font.name = 'Arial'
        normal.font.size = Pt(11)

        # Nomor surat pada dokumen Word sengaja dibuat manual/editable.
        # Pengelola/Biro Umum mengisi nomor final langsung pada file Word sebelum PDF final diupload.
        nomor = '........................................'
        tanggal = obj.tanggal_nota_permohonan_psp or obj.tanggal_permohonan or timezone.now().date()
        nama_pejabat, nip_pejabat, jabatan_pejabat = _pejabat_psp(obj)
        satuan, jumlah, nilai, judul, tiket = _psp_summary(obj)

        if jenis == 'sk':
            # Konsep SK Penetapan PSP untuk role Sekjen. Format mengikuti contoh SK PSP:
            # halaman keputusan tanpa lampiran daftar barang. Nomor SK dibuat manual/editable.
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt

            section = doc.sections[0]
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.4)
            section.right_margin = Cm(2.2)

            sk_nomor = '....../HUK/........'
            jenis_objek = (obj.get_jenis_barang_display() or 'Barang Milik Negara').upper()
            satuan_upper = (satuan or '-').upper()
            jumlah_text = f'{jumlah} ({jumlah}) unit' if jumlah else '........ unit'

            _docx_p(doc, 'KEMENTERIAN SOSIAL REPUBLIK INDONESIA', bold=True, align='center', size=12, space_after=10)
            _docx_p(doc, 'KEPUTUSAN MENTERI SOSIAL REPUBLIK INDONESIA', bold=True, align='center', size=11, space_after=0)
            _docx_p(doc, f'NOMOR {sk_nomor}', bold=True, align='center', size=11, space_after=8)
            _docx_p(doc, 'TENTANG', bold=True, align='center', size=11, space_after=4)
            _docx_p(doc, f'PENETAPAN STATUS PENGGUNAAN BARANG MILIK NEGARA SELAIN TANAH DAN/ATAU BANGUNAN BERUPA {jenis_objek} PADA {satuan_upper}', bold=True, align='center', size=10, space_after=12)
            _docx_p(doc, 'MENTERI SOSIAL REPUBLIK INDONESIA,', bold=True, align='center', size=10, space_after=12)

            _docx_meta_table(doc, [
                ('Menimbang', f'a. bahwa untuk efisiensi, efektivitas, dan tertib administrasi pengelolaan terhadap barang milik negara yang masih layak pakai untuk kepentingan dinas, perlu melakukan penetapan status penggunaan barang milik negara selain tanah dan/atau bangunan berupa {jumlah_text} {obj.get_jenis_barang_display() or "Barang Milik Negara"} pada {satuan};'),
                ('', 'b. bahwa berdasarkan pertimbangan sebagaimana dimaksud dalam huruf a, perlu menetapkan Keputusan Menteri Sosial tentang Penetapan Status Penggunaan Barang Milik Negara Selain Tanah dan/atau Bangunan;'),
            ])
            _docx_meta_table(doc, [
                ('Mengingat', '1. Undang-Undang Nomor 1 Tahun 2004 tentang Perbendaharaan Negara;'),
                ('', '2. Peraturan Pemerintah Nomor 27 Tahun 2014 tentang Pengelolaan Barang Milik Negara/Daerah sebagaimana telah diubah dengan Peraturan Pemerintah Nomor 28 Tahun 2020;'),
                ('', '3. Peraturan Presiden Nomor 162 Tahun 2024 tentang Kementerian Sosial;'),
                ('', '4. Peraturan Menteri Sosial tentang Tata Cara Pelaksanaan Penggunaan, Pemanfaatan, Penghapusan, dan Pemindahtanganan Barang Milik Negara di Lingkungan Kementerian Sosial;'),
                ('', '5. Peraturan Menteri Keuangan Nomor 4/PMK.06/2015 tentang Pendelegasian Kewenangan dan Tanggung Jawab Tertentu dari Pengelola Barang kepada Pengguna Barang;'),
                ('', '6. Peraturan Menteri Keuangan Nomor 40 Tahun 2024 tentang Tata Cara Penggunaan Barang Milik Negara;'),
                ('', '7. Keputusan Menteri Sosial tentang Pendelegasian Kewenangan dari Menteri Sosial kepada Sekretaris Jenderal terhadap Pengelolaan Barang Milik Negara;'),
            ])
            _docx_meta_table(doc, [
                ('Memperhatikan', '1. Surat Pernyataan Formil dan Materiil;'),
                ('', '2. Surat Keterangan Kebenaran Dokumen Digital; dan'),
                ('', '3. Nota Dinas permohonan Penetapan Status Penggunaan Barang Milik Negara.'),
            ])

            _docx_p(doc, 'MEMUTUSKAN:', bold=True, align='center', size=11, space_after=8)
            _docx_meta_table(doc, [
                ('Menetapkan', f'KEPUTUSAN MENTERI SOSIAL TENTANG PENETAPAN STATUS PENGGUNAAN BARANG MILIK NEGARA SELAIN TANAH DAN/ATAU BANGUNAN BERUPA {jenis_objek} PADA {satuan_upper}.'),
                ('KESATU', f'Menetapkan status penggunaan barang milik negara berupa {jumlah_text} {obj.get_jenis_barang_display() or "Barang Milik Negara"} pada {satuan}.'),
                ('KEDUA', f'Nilai perolehan barang milik negara sebagaimana dimaksud dalam Diktum KESATU seluruhnya sebesar {_rupiah_plain(nilai)}.'),
                ('KETIGA', 'Barang milik negara sebagaimana dimaksud dalam Diktum KESATU dicatat dalam daftar barang kuasa pengguna pada kuasa pengguna barang.'),
                ('KEEMPAT', 'Pengguna barang dapat melakukan pemanfaatan atau pemindahtanganan kepada pihak lain setelah mendapatkan persetujuan pengelola barang sesuai dengan ketentuan peraturan perundang-undangan.'),
                ('KELIMA', 'Pengguna barang mempunyai wewenang melakukan pengawasan dan pengendalian atas barang milik negara sebagaimana dimaksud dalam Diktum KESATU.'),
                ('KEENAM', f'Segala pembiayaan pengamanan dan pemeliharaan barang milik negara yang digunakan menjadi tanggung jawab {satuan}.'),
                ('KETUJUH', 'Keputusan Menteri ini mulai berlaku pada tanggal ditetapkan, dengan ketentuan apabila di kemudian hari terdapat kekeliruan dalam penetapannya akan diperbaiki sebagaimana mestinya.'),
            ])

            table = doc.add_table(rows=1, cols=2)
            _docx_set_table_borders_none(table)
            table.columns[0].width = Cm(9.0)
            table.columns[1].width = Cm(7.0)
            cell = table.cell(0, 1)
            for line in [
                'Ditetapkan di Jakarta',
                'pada tanggal ........................',
                'a.n. MENTERI SOSIAL REPUBLIK INDONESIA',
                'SEKRETARIS JENDERAL,',
                '', '', '',
                '........................................',
            ]:
                par = cell.add_paragraph(line)
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in par.runs:
                    _docx_set_font(r, size=10, bold=('SEKRETARIS JENDERAL' in line or 'MENTERI SOSIAL' in line))

            doc.add_page_break()
            _docx_p(doc, 'Salinan Keputusan Menteri ini disampaikan kepada:', size=10)
            for line in [
                '1. Menteri Sosial;',
                '2. Inspektur Jenderal, Kementerian Sosial;',
                '3. Kepala Kantor Pelayanan Kekayaan Negara dan Lelang setempat;',
                '4. Kepala Biro Umum, Kementerian Sosial.',
            ]:
                _docx_p(doc, line, size=10)
            filename = f'sk_penetapan_psp_sekjen_{obj.pk}.docx'
            return _docx_response(doc, filename)

        else:
            _docx_add_kop_psp(doc)

        if jenis == 'nota':
            _docx_p(doc, 'NOTA DINAS', bold=True, align='center', size=13, space_after=0)
            _docx_p(doc, f'NOMOR: {nomor}', bold=True, align='center', size=10, space_after=8)
            _docx_meta_table(doc, [
                ('Yth', 'Sekretaris Jenderal'),
                ('Dari', 'Kepala Biro Umum'),
                ('Hal', judul),
                ('Lampiran', '1 Berkas'),
                ('Sifat', 'Segera'),
                ('Tanggal', _psp_doc_date(tanggal)),
            ])
            _docx_p(doc, '', space_after=2)
            _docx_p(doc, '1. Dasar:', size=10)
            for line in [
                'a. Peraturan Pemerintah Nomor 28 Tahun 2020 tentang Pengelolaan Barang Milik Negara/Daerah;',
                'b. Peraturan Menteri Keuangan Nomor 40 Tahun 2024 tentang Tata Cara Penggunaan Barang Milik Negara;',
                'c. Peraturan Menteri Keuangan Nomor 4/PMK.06/2015 tentang Pendelegasian Kewenangan dan Tanggung Jawab Tertentu Dari Pengelola Barang Kepada Pengguna Barang.',
            ]:
                _docx_p(doc, line, size=10, align='justify')
            _docx_p(doc, f'2. Dalam rangka tertib administrasi Pengelolaan dan Penatausahaan Barang Milik Negara pada Kementerian Sosial, dengan ini kami mengajukan Permohonan Penetapan Status Penggunaan Barang Milik Negara pada {satuan} dengan jumlah {jumlah} unit dan total nilai perolehan {_rupiah_plain(nilai)} dengan nomor tiket terdaftar pada SIMAN V2: {tiket}.', size=10, align='justify')
            _docx_p(doc, '3. Sebagai bahan pertimbangan dalam menetapkan Status Penggunaan Barang Milik Negara tersebut, bersama ini kami sertakan kelengkapan data/dokumen sebagai berikut:', size=10, align='justify')
            for line in ['a. Surat Pernyataan Formil dan Materiil', 'b. Surat Keterangan Kebenaran Dokumen Digital', 'c. Laporan Sub-Sub Kelompok', 'd. Laporan Kondisi Barang']:
                _docx_p(doc, '   ' + line, size=10)
            _docx_p(doc, 'Demikian disampaikan, atas perhatian dan kerjasamanya disampaikan terima kasih.', size=10, align='justify')
            _docx_signature(doc, jabatan_pejabat, nama_pejabat, nip_pejabat, tanggal, meterai=False)
            filename = f'nota_dinas_psp_{obj.pk}.docx'

        elif jenis == 'keterangan':
            _docx_p(doc, 'SURAT KETERANGAN', bold=True, align='center', size=13, space_after=0)
            _docx_p(doc, f'NOMOR: {nomor}', bold=True, align='center', size=10, space_after=8)
            _docx_p(doc, 'Yang bertanda tangan di bawah ini:', size=10)
            _docx_identitas_table(doc, nama_pejabat, nip_pejabat, jabatan_pejabat)
            _docx_p(doc, f'Dengan ini menerangkan bahwa semua dokumen arsip digital dalam rangka permohonan penetapan status Barang Milik Negara berupa {judul} pada {satuan} dengan nilai perolehan {_rupiah_plain(nilai)} adalah sesuai dengan aslinya.', size=10, align='justify')
            _docx_p(doc, 'Demikian surat keterangan ini kami buat dengan sebenar-benarnya dalam rangka tertib administrasi pengelolaan BMN.', size=10, align='justify')
            _docx_signature(doc, jabatan_pejabat, nama_pejabat, nip_pejabat, tanggal, meterai=False)
            filename = f'surat_keterangan_digital_psp_{obj.pk}.docx'

        else:
            _docx_p(doc, 'SURAT PERNYATAAN', bold=True, align='center', size=13, space_after=0)
            _docx_p(doc, f'NOMOR: {nomor}', bold=True, align='center', size=10, space_after=8)
            _docx_p(doc, 'Yang bertanda tangan di bawah ini,', size=10)
            _docx_identitas_table(doc, nama_pejabat, nip_pejabat, jabatan_pejabat)
            _docx_p(doc, 'Dengan ini menyatakan bahwa:', size=10)
            _docx_p(doc, f'1. Bertanggung jawab atas kebenaran formil dan materiil Barang Milik Negara yang diajukan permohonan penetapan status penggunaan sebanyak {jumlah} unit dengan total nilai perolehan {_rupiah_plain(nilai)}, dikuasai dan digunakan untuk penyelenggaraan tugas dan fungsi pada satuan kerja {satuan}.', size=10, align='justify')
            _docx_p(doc, '2. Barang Milik Negara yang diajukan untuk Penetapan Status Penggunaan digunakan dalam rangka pengamanan BMN.', size=10, align='justify')
            _docx_p(doc, 'Demikian pernyataan ini kami buat dengan sebenarnya dalam rangka tertib administrasi pengelolaan Barang Milik Negara.', size=10, align='justify')
            _docx_p(doc, 'Catatan: Letakkan e-Meterai elektronik berdampingan dengan area TTE/QR BSrE dan jangan menimpa QR TTE.', size=9)
            _docx_signature(doc, jabatan_pejabat, nama_pejabat, nip_pejabat, tanggal, meterai=True)
            filename = f'surat_pernyataan_formil_materiil_psp_{obj.pk}.docx'

        return _docx_response(doc, filename)


class DownloadTemplateBarangPSPView(LoginRequiredMixin, View):
    def get(self, request, *args, **kwargs):
        wb = Workbook()
        ws = wb.active
        ws.title = 'Lampiran PSP'
        headers = ['no', 'kode_satuan_kerja', 'nama_satuan_kerja', 'kode_barang', 'nup', 'nama_barang', 'tipe_barang', 'tahun_perolehan', 'kuantitas', 'nilai_perolehan', 'kondisi_barang', 'keterangan']
        ws.append(headers)
        ws.append([1, '027010199440121006KP', 'Sekretariat Jenderal Sekolah Rakyat', '3100102002', '9471', 'Lap Top', 'Acer Travelmate P214 Core 5 8GB/256GB', '20-11-2025', 1, 9183030, 'Baik', 'Contoh data'])
        for col in ws.columns:
            ws.column_dimensions[col[0].column_letter].width = min(max(len(str(col[0].value or '')) + 4, 14), 34)
        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)
        response = HttpResponse(bio.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = 'attachment; filename="template_import_barang_psp.xlsx"'
        return response


class ExportLampiranPSPPDFView(PermohonanPSPAccessMixin, View):
    def get(self, request, *args, **kwargs):
        obj = get_object_or_404(self.get_scoped_queryset(), pk=kwargs['pk'])
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), leftMargin=0.8*cm, rightMargin=0.8*cm, topMargin=0.8*cm, bottomMargin=0.8*cm)
        styles = getSampleStyleSheet()
        story = [
            Paragraph('LAMPIRAN DAFTAR BARANG MILIK NEGARA YANG DITETAPKAN STATUS PENGGUNAANNYA', styles['Title']),
            Paragraph(obj.judul_paket or obj.get_jenis_barang_display(), styles['Normal']),
            Paragraph(f'Nomor Permohonan: {obj.nomor_permohonan} | Tiket SIMAN: {obj.nomor_tiket_siman or "-"}', styles['Normal']),
            Spacer(1, 0.3*cm),
        ]
        data = [['NO', 'KODE SATKER', 'NAMA SATKER', 'KODE BARANG', 'NUP', 'NAMA BARANG', 'TIPE', 'TAHUN', 'QTY', 'NILAI', 'KONDISI']]
        for item in obj.detail_barang.all():
            data.append([
                item.nomor_urut, item.kode_satuan_kerja or '-', item.nama_satuan_kerja or '-', item.kode_barang,
                item.nup, item.nama_barang, item.tipe_barang or '-', item.tahun_perolehan or '-', item.kuantitas,
                f'{item.nilai_perolehan:,.0f}'.replace(',', '.'), item.get_kondisi_barang_display(),
            ])
        if len(data) == 1:
            data.append(['-', '-', '-', '-', '-', 'Belum ada detail barang', '-', '-', '-', '-', '-'])
        table = Table(data, repeatRows=1, colWidths=[1*cm, 2.7*cm, 3.5*cm, 2.4*cm, 1.5*cm, 3.5*cm, 4.2*cm, 2*cm, 1.1*cm, 2.2*cm, 1.7*cm])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('FONTSIZE', (0, 0), (-1, -1), 6),
            ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph(f'Total barang: {obj.jumlah_barang} unit | Total nilai: Rp{obj.total_nilai_barang:,.0f}'.replace(',', '.'), styles['Normal']))
        story.append(Paragraph('Catatan: Dokumen final wajib ditandatangani menggunakan TTE BSrE. Apabila membutuhkan meterai, gunakan e-Meterai resmi dan tempatkan berdampingan dengan area TTE, tidak menimpa QR/visual e-Meterai.', styles['Normal']))
        doc.build(story)
        buffer.seek(0)
        filename = f'lampiran_psp_{obj.pk}.pdf'
        return FileResponse(buffer, as_attachment=True, filename=filename)


# =============================================================
# Export daftar Permohonan PSP BMN (PDF, Excel, CSV)
# =============================================================
def _psp_columns():
    return [
        ('No', '__no__'),
        ('Nomor Permohonan', 'nomor_permohonan'),
        ('Tanggal Permohonan', 'tanggal_permohonan'),
        ('Nomor Nota ke Sekjen', 'nomor_nota_permohonan_psp'),
        ('Tiket SIMAN', 'nomor_tiket_siman'),
        ('Unit Kerja', 'unit_kerja__nama_unit'),
        ('Pemohon', 'pemohon__nama'),
        ('NIP Pemohon', 'pemohon__nip'),
        ('Judul Paket', 'judul_paket'),
        ('Jenis Barang', 'display:jenis_barang'),
        ('Nama Barang', 'nama_barang'),
        ('Kode Barang', 'kode_barang'),
        ('NUP', 'nup'),
        ('Nilai PSP', 'nilai_psp'),
        ('Kategori Nilai', lambda o: getattr(o, 'kategori_nilai_display', '') or ''),
        ('Status', 'display:status'),
        ('Nomor SK PSP', 'nomor_sk_psp'),
        ('Tanggal SK PSP', 'tanggal_sk_psp'),
        ('Catatan Biro Umum', 'catatan_biro_umum'),
        ('Catatan Unit', 'catatan_unit'),
    ]


def _psp_queryset_for_mode(request, mode):
    view = PermohonanPSPListView()
    view.request = request
    view.mode = mode or 'permohonan'
    if mode == 'verifikasi':
        view.__class__ = VerifikasiPSPListView
    elif mode == 'persetujuan_sekjen':
        view.__class__ = PersetujuanSekjenPSPListView
    qs = view.get_base_queryset_for_mode()
    return apply_search_filter(qs, request, PermohonanPSPListView.search_fields)


@login_required
def export_psp(request, fmt):
    mode = (request.GET.get('mode') or 'permohonan').strip()
    qs = _psp_queryset_for_mode(request, mode)
    title_map = {
        'verifikasi': 'Verifikasi Usulan PSP BMN',
        'persetujuan_sekjen': 'Persetujuan/Penetapan PSP BMN',
    }
    filename = f'transaksi_psp_bmn_{mode}'
    return export_queryset(request, qs, fmt, filename, title_map.get(mode, 'Permohonan PSP BMN'), _psp_columns(), order_by=['-tanggal_permohonan', '-id'])
